from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from backend.document_loader import extract_text
from backend.domain.chunker import DocumentObject, detect_repeated_noise_lines, infer_document_type, infer_source_year, split_text_to_objects, source_hash


@dataclass(slots=True)
class ProcessedDocument:
    source_name: str
    source_hash: str
    objects: list[DocumentObject]

    @property
    def total_chars(self) -> int:
        return sum(len(obj.text) for obj in self.objects)


def _safe_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\n", " ").strip()


def _markdown_table(rows: list[list[Any]]) -> str:
    cleaned = [[_safe_cell(cell) for cell in row] for row in rows]
    cleaned = [row for row in cleaned if any(cell for cell in row)]
    if not cleaned:
        return ""
    width = max(len(row) for row in cleaned)
    padded = [row + [""] * (width - len(row)) for row in cleaned]
    header = padded[0]
    sep = ["---"] * width
    body = padded[1:]
    return "\n".join(
        ["| " + " | ".join(header) + " |", "| " + " | ".join(sep) + " |"]
        + ["| " + " | ".join(row) + " |" for row in body]
    )


def _extract_docx_objects(path: Path, source_name: str, base_metadata: dict[str, Any]) -> list[DocumentObject]:
    from docx import Document

    doc = Document(str(path))
    text_parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            text_parts.append(p.text.strip())

    objects = split_text_to_objects("\n\n".join(text_parts), source_name=source_name, base_metadata=base_metadata)
    for idx, table in enumerate(doc.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        table_text = _markdown_table(rows)
        if not table_text:
            continue
        objects.extend(
            split_text_to_objects(
                table_text,
                source_name=source_name,
                base_metadata={**base_metadata, "table_index": idx, "document_kind": "table_from_docx"},
                max_chars=6000,
            )
        )
    return objects


def _extract_xlsx_objects(path: Path, source_name: str, base_metadata: dict[str, Any]) -> list[DocumentObject]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    objects: list[DocumentObject] = []
    for ws in wb.worksheets:
        rows: list[list[Any]] = []
        for row in ws.iter_rows(values_only=True):
            rows.append(list(row))
        table_text = _markdown_table(rows)
        if not table_text:
            continue
        objects.extend(
            split_text_to_objects(
                table_text,
                source_name=source_name,
                base_metadata={**base_metadata, "sheet": ws.title, "document_kind": "spreadsheet"},
                max_chars=6000,
            )
        )
    return objects


def _extract_csv_objects(path: Path, source_name: str, base_metadata: dict[str, Any]) -> list[DocumentObject]:
    raw = path.read_bytes()
    text = None
    for enc in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = raw.decode("utf-8", errors="ignore")
    
    try:
        dialect = csv.Sniffer().sniff(text[:4096]) if text.strip() else csv.excel
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect=dialect))
    table_text = _markdown_table(rows)
    return split_text_to_objects(
        table_text or text,
        source_name=source_name,
        base_metadata={**base_metadata, "document_kind": "csv_table"},
        max_chars=6000,
    )


def _extract_pdf_objects(path: Path, source_name: str, base_metadata: dict[str, Any]) -> list[DocumentObject]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        page_texts: list[tuple[int, str]] = []
        for page_idx, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                page_texts.append((page_idx, text))
        repeated = detect_repeated_noise_lines([text for _idx, text in page_texts], min_repeats=3)
        objects: list[DocumentObject] = []
        for page_idx, text in page_texts:
            objects.extend(
                split_text_to_objects(
                    text,
                    source_name=source_name,
                    base_metadata={**base_metadata, "page": page_idx, "repeated_noise_lines": list(repeated)},
                )
            )
        if objects:
            return objects
    except Exception:
        pass
    text = extract_text(path)
    repeated = detect_repeated_noise_lines([text], min_repeats=3)
    return split_text_to_objects(text, source_name=source_name, base_metadata={**base_metadata, "repeated_noise_lines": list(repeated)})


def process_document(path: Path, *, original_name: str | None = None) -> ProcessedDocument:
    source_name = original_name or path.name
    file_hash = source_hash(path)
    suffix = path.suffix.lower()
    preliminary_type = infer_document_type(source_name, "", {"extension": suffix})
    base_metadata = {
        "source_hash": file_hash,
        "source": source_name,
        "extension": suffix,
        "document_type": preliminary_type,
        "source_year": infer_source_year(source_name),
    }

    if suffix == ".docx":
        objects = _extract_docx_objects(path, source_name, base_metadata)
    elif suffix == ".xlsx":
        objects = _extract_xlsx_objects(path, source_name, base_metadata)
    elif suffix in {".csv", ".tsv"}:
        objects = _extract_csv_objects(path, source_name, base_metadata)
    elif suffix == ".pdf":
        objects = _extract_pdf_objects(path, source_name, base_metadata)
    else:
        text = extract_text(path)
        objects = split_text_to_objects(text, source_name=source_name, base_metadata=base_metadata)

    if not objects:
        raise ValueError("Документ не содержит извлекаемых объектов")
    return ProcessedDocument(source_name=source_name, source_hash=file_hash, objects=objects)
